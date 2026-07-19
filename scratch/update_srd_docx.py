import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

doc_path = r"folder_template/Template2_SRD Document.docx"
if not os.path.exists(doc_path):
    print("File not found:", doc_path)
    sys.exit(1)

doc = docx.Document(doc_path)
print("Loaded document successfully. Total paragraphs:", len(doc.paragraphs))

# 1. Clean Placeholders in Text Paragraphs
replacements = {
    "[This part shows the system screens and the relationship among screens. You can draw the Screens Flow for the system in the form of diagram as below]":
        "The screen flow defines the navigation paths and user interface transitions across different roles (Admin, Horse Owner, Jockey, Referee, Spectator) in the Horse Racing Tournament Management System.",
    "[Provide the tables relationship like example below  following MySQL database naming convention]":
        "The ERD diagram below outlines the relational database structure, foreign key references, and entity relationships for the Horse Racing Tournament System.",
    "[Give some lines about the table here>>":
        "The following section details the schema, data types, primary/foreign key constraints, and field descriptions for all 17 database tables in the system.",
    "[Table fields, in the form of table format as below]":
        "",
    "[Table] User": "Table 1: User (System Accounts & Profiles)",
    "[Table] Race": "Table 2: Race (Individual Race Events)",
    "[Table] Horse": "Table 3: Horse (Horse Profiles & Statistics)",
    "[Table] RaceEntry": "Table 4: RaceEntry (Race Participation & Results)"
}

for p in doc.paragraphs:
    for old_t, new_t in replacements.items():
        if old_t in p.text:
            p.text = p.text.replace(old_t, new_t)

# Fix any leftover '>>' in State transition diagram section
for p in doc.paragraphs:
    if p.text.strip() == ">>":
        p.text = ""

print("Cleaned up template placeholder texts.")

# 2. Add Missing Functional Requirements in Section II
# Let's locate the headings for Horse Owner, Jockey, Spectator, Admin to insert new FEs
def add_fe_block(doc, target_p_index, title, trigger, desc, logic):
    p_title = doc.paragraphs[target_p_index].insert_paragraph_before(title, style='Heading 3')
    
    p_trig = doc.paragraphs[target_p_index].insert_paragraph_before(f"• Function trigger: {trigger}")
    p_desc = doc.paragraphs[target_p_index].insert_paragraph_before(f"• Function description: {desc}")
    p_logic = doc.paragraphs[target_p_index].insert_paragraph_before(f"• Business logic & validations: {logic}")
    doc.paragraphs[target_p_index].insert_paragraph_before("")

# Find indices for insertions:
# After FE-08 Horse Owner Management (before Jockey Participation Management)
jockey_section_idx = None
spectator_section_idx = None
admin_section_idx = None
db_section_idx = None

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == "2. Jockey Participation Management":
        jockey_section_idx = i
    elif t == "3. Race Officiating System":
        referee_section_idx = i
    elif t == "4. Spectator Management System":
        spectator_section_idx = i
    elif t == "5. Race Management & Administration System":
        admin_section_idx = i
    elif t == "III. Database Design":
        db_section_idx = i

print(f"Section indices: Jockey={jockey_section_idx}, Referee={referee_section_idx}, Spectator={spectator_section_idx}, Admin={admin_section_idx}, DB={db_section_idx}")

# Insert new FEs
# Under 1. Horse Owner Management: FE-09 & FE-10
add_fe_block(doc, jockey_section_idx, 
             "FE-09: Register Horses for Race Meetings",
             "Owner Dashboard -> 'Race Meeting Registration'.",
             "Horse Owner registers eligible horses for an upcoming scheduled Race Meeting.",
             "Verifies horse status (must be ACTIVE) and ensures horse is not already registered for the same meeting. Creates a record in HorseRaceMeetingRegistration with PENDING status.")

add_fe_block(doc, jockey_section_idx, 
             "FE-10: Submit Horse Retirement Request",
             "Owner Dashboard -> My Horses -> 'Request Retirement'.",
             "Horse Owner submits a formal request with a detailed reason to retire an active horse from competition.",
             "Creates a record in HorseRetirementRequest with status PENDING. Horse cannot be entered into new races while retirement request is pending.")

# Under 2. Jockey Participation Management: FE-07
add_fe_block(doc, referee_section_idx, 
             "FE-07: Register Availability for Race Meetings",
             "Jockey Dashboard -> 'Race Meeting Registration'.",
             "Jockey registers availability to ride in races scheduled for an upcoming Race Meeting.",
             "Verifies jockey account status is ACTIVE. Creates a record in JockeyRaceMeetingRegistration with PENDING status for Admin review.")

# Under 4. Spectator Management System: FE-04
add_fe_block(doc, admin_section_idx, 
             "FE-04: AI Race Outcome Predictor & Smart Assistant",
             "Landing Portal -> 'AI Race Assistant / Predictor'.",
             "Spectator interacts with the AI Assistant to query race stats and obtain AI-driven predictions for upcoming race outcomes.",
             "Integrates with Python backend AI engine (chatbot.py, predictor.py) utilizing multi-key API rotation (Gemini & Groq) for real-time statistical predictions based on horse ratings, jockey win rates, and track conditions.")

# Under 5. Race Management & Administration System: FE-09, FE-10, FE-11
add_fe_block(doc, db_section_idx, 
             "FE-09: Manage Race Meeting Registrations",
             "Admin Dashboard -> 'Meeting Registrations' -> Approve / Reject.",
             "Admin reviews and approves or rejects registrations from Horse Owners, Jockeys, and Horses for upcoming Race Meetings.",
             "Updates registration status in OwnerRaceMeetingRegistration, JockeyRaceMeetingRegistration, and HorseRaceMeetingRegistration. Only approved participants can be invited or declared for races in that meeting.")

add_fe_block(doc, db_section_idx, 
             "FE-10: Process Horse Retirement Requests",
             "Admin Dashboard -> 'Retirement Requests'.",
             "Admin reviews pending horse retirement requests submitted by owners, adds administrative remarks, and approves or rejects them.",
             "Approving sets HorseRetirementRequest status to APPROVED and updates Horse status to RETIRED. Rejection returns horse status to ACTIVE.")

add_fe_block(doc, db_section_idx, 
             "FE-11: Configure Live Stream Settings",
             "Admin Dashboard -> 'Live Settings'.",
             "Admin configures live streaming options and sets the YouTube Live URL for scheduled races.",
             "Validates URL format and updates youtube_live_url in Race table. Embedded player on spectator dashboard updates automatically.")

print("Added missing Functional Requirements.")

# 3. Add 13 Missing Database Tables in Section III.2
missing_tables = [
    {
        "title": "Table 5: Role (User Roles & Permissions)",
        "desc": "Stores system security roles (Admin, Owner, Jockey, Spectator, Referee).",
        "rows": [
            ["#", "Field name", "Type", "Size", "Unique", "Not Null", "PK/FK", "Notes"],
            ["1", "id", "INT", "4", "Yes", "Yes", "PK", "Identity (1,1) auto-increment"],
            ["2", "role_name", "VARCHAR", "50", "Yes", "Yes", "-", "Role name ('Admin', 'Owner', 'Jockey', 'Spectator', 'Referee')"]
        ]
    },
    {
        "title": "Table 6: Season (Championship Seasons)",
        "desc": "Stores tournament seasons and date ranges.",
        "rows": [
            ["#", "Field name", "Type", "Size", "Unique", "Not Null", "PK/FK", "Notes"],
            ["1", "id", "INT", "4", "Yes", "Yes", "PK", "Identity (1,1) auto-increment"],
            ["2", "name", "VARCHAR", "200", "No", "Yes", "-", "Tournament season title"],
            ["3", "start_date", "DATE", "3", "No", "Yes", "-", "Season start date"],
            ["4", "end_date", "DATE", "3", "No", "Yes", "-", "Season end date"],
            ["5", "status", "VARCHAR", "20", "No", "Yes", "-", "Status ('ACTIVE', 'CLOSED')"]
        ]
    },
    {
        "title": "Table 7: SeasonClassRule (Class Brackets & Rules)",
        "desc": "Defines rating brackets and prize ranges for race classes (Class 1 to Class 5).",
        "rows": [
            ["#", "Field name", "Type", "Size", "Unique", "Not Null", "PK/FK", "Notes"],
            ["1", "id", "INT", "4", "Yes", "Yes", "PK", "Identity (1,1) auto-increment"],
            ["2", "season_id", "INT", "4", "No", "Yes", "FK", "References Season(id)"],
            ["3", "class_level", "VARCHAR", "50", "No", "Yes", "-", "Class level ('Class 1', 'Class 2', etc.)"],
            ["4", "class_name", "VARCHAR", "100", "No", "No", "-", "Descriptive class bracket name"],
            ["5", "min_rating", "INT", "4", "No", "No", "-", "Minimum rating required"],
            ["6", "max_rating", "INT", "4", "No", "No", "-", "Maximum rating allowed"],
            ["7", "min_prize", "DECIMAL", "18,2", "No", "No", "-", "Minimum prize allocation"],
            ["8", "max_prize", "DECIMAL", "18,2", "No", "No", "-", "Maximum prize allocation"]
        ]
    },
    {
        "title": "Table 8: RaceMeeting (Race Meeting Events)",
        "desc": "Defines scheduled race day events within a season.",
        "rows": [
            ["#", "Field name", "Type", "Size", "Unique", "Not Null", "PK/FK", "Notes"],
            ["1", "id", "INT", "4", "Yes", "Yes", "PK", "Identity (1,1) auto-increment"],
            ["2", "season_id", "INT", "4", "No", "Yes", "FK", "References Season(id)"],
            ["3", "name", "VARCHAR", "200", "No", "Yes", "-", "Race meeting title"],
            ["4", "start_date", "DATETIME", "8", "No", "Yes", "-", "Meeting event date & time"],
            ["5", "venue", "VARCHAR", "150", "No", "Yes", "-", "Track venue location"],
            ["6", "total_budget", "DECIMAL", "18,2", "No", "Yes", "-", "Total prize budget"]
        ]
    },
    {
        "title": "Table 9: OwnerRaceMeetingRegistration (Owner Event Registration)",
        "desc": "Tracks horse owner registrations for specific race meetings.",
        "rows": [
            ["#", "Field name", "Type", "Size", "Unique", "Not Null", "PK/FK", "Notes"],
            ["1", "id", "INT", "4", "Yes", "Yes", "PK", "Identity (1,1) auto-increment"],
            ["2", "race_meeting_id", "INT", "4", "No", "Yes", "FK", "References RaceMeeting(id)"],
            ["3", "owner_id", "INT", "4", "No", "Yes", "FK", "References User(id)"],
            ["4", "status", "VARCHAR", "20", "No", "Yes", "-", "Status ('PENDING', 'APPROVED', 'REJECTED')"],
            ["5", "registered_at", "DATETIME", "8", "No", "Yes", "-", "Registration timestamp"]
        ]
    },
    {
        "title": "Table 10: JockeyRaceMeetingRegistration (Jockey Event Registration)",
        "desc": "Tracks jockey availability registrations for race meetings.",
        "rows": [
            ["#", "Field name", "Type", "Size", "Unique", "Not Null", "PK/FK", "Notes"],
            ["1", "id", "INT", "4", "Yes", "Yes", "PK", "Identity (1,1) auto-increment"],
            ["2", "race_meeting_id", "INT", "4", "No", "Yes", "FK", "References RaceMeeting(id)"],
            ["3", "jockey_id", "INT", "4", "No", "Yes", "FK", "References User(id)"],
            ["4", "status", "VARCHAR", "20", "No", "Yes", "-", "Status ('PENDING', 'APPROVED', 'REJECTED')"],
            ["5", "registered_at", "DATETIME", "8", "No", "Yes", "-", "Registration timestamp"]
        ]
    },
    {
        "title": "Table 11: HorseRaceMeetingRegistration (Horse Event Registration)",
        "desc": "Tracks horse registrations for race meetings.",
        "rows": [
            ["#", "Field name", "Type", "Size", "Unique", "Not Null", "PK/FK", "Notes"],
            ["1", "id", "INT", "4", "Yes", "Yes", "PK", "Identity (1,1) auto-increment"],
            ["2", "race_meeting_id", "INT", "4", "No", "Yes", "FK", "References RaceMeeting(id)"],
            ["3", "horse_id", "INT", "4", "No", "Yes", "FK", "References Horse(id)"],
            ["4", "status", "VARCHAR", "20", "No", "Yes", "-", "Status ('PENDING', 'APPROVED', 'REJECTED')"],
            ["5", "registered_at", "DATETIME", "8", "No", "Yes", "-", "Registration timestamp"]
        ]
    },
    {
        "title": "Table 12: RaceInvitation (Jockey Invitations)",
        "desc": "Stores invitations sent by horse owners to jockeys for specific races.",
        "rows": [
            ["#", "Field name", "Type", "Size", "Unique", "Not Null", "PK/FK", "Notes"],
            ["1", "id", "INT", "4", "Yes", "Yes", "PK", "Identity (1,1) auto-increment"],
            ["2", "race_id", "INT", "4", "No", "Yes", "FK", "References Race(id)"],
            ["3", "horse_id", "INT", "4", "No", "Yes", "FK", "References Horse(id)"],
            ["4", "owner_id", "INT", "4", "No", "Yes", "FK", "References User(id)"],
            ["5", "jockey_id", "INT", "4", "No", "Yes", "FK", "References User(id)"],
            ["6", "status", "VARCHAR", "20", "No", "Yes", "-", "Status ('PENDING', 'ACCEPTED', 'REJECTED', 'EXPIRED')"]
        ]
    },
    {
        "title": "Table 13: RaceReferee (Referee Race Assignment)",
        "desc": "Maps officiating referees to assigned race events.",
        "rows": [
            ["#", "Field name", "Type", "Size", "Unique", "Not Null", "PK/FK", "Notes"],
            ["1", "id", "INT", "4", "Yes", "Yes", "PK", "Identity (1,1) auto-increment"],
            ["2", "race_id", "INT", "4", "No", "Yes", "FK", "References Race(id)"],
            ["3", "referee_id", "INT", "4", "No", "Yes", "FK", "References User(id)"]
        ]
    },
    {
        "title": "Table 14: Violation (Race Rule Violations)",
        "desc": "Records rule infractions and penalties issued by referees during races.",
        "rows": [
            ["#", "Field name", "Type", "Size", "Unique", "Not Null", "PK/FK", "Notes"],
            ["1", "id", "INT", "4", "Yes", "Yes", "PK", "Identity (1,1) auto-increment"],
            ["2", "race_id", "INT", "4", "No", "Yes", "FK", "References Race(id)"],
            ["3", "horse_id", "INT", "4", "No", "Yes", "FK", "References Horse(id)"],
            ["4", "jockey_id", "INT", "4", "No", "Yes", "FK", "References User(id)"],
            ["5", "referee_id", "INT", "4", "No", "Yes", "FK", "References User(id)"],
            ["6", "description", "VARCHAR", "500", "No", "Yes", "-", "Infraction details"],
            ["7", "penalty", "VARCHAR", "200", "No", "Yes", "-", "Penalty or fine imposed"],
            ["8", "status", "VARCHAR", "30", "No", "Yes", "-", "Status ('PENDING', 'RESOLVED')"]
        ]
    },
    {
        "title": "Table 15: ChatMessage (Live Stream Chat)",
        "desc": "Stores live spectator chat messages per race.",
        "rows": [
            ["#", "Field name", "Type", "Size", "Unique", "Not Null", "PK/FK", "Notes"],
            ["1", "id", "INT", "4", "Yes", "Yes", "PK", "Identity (1,1) auto-increment"],
            ["2", "race_id", "INT", "4", "No", "Yes", "FK", "References Race(id)"],
            ["3", "username", "VARCHAR", "100", "No", "Yes", "-", "Spectator username"],
            ["4", "message_text", "NVARCHAR", "MAX", "No", "Yes", "-", "Chat message text"],
            ["5", "sent_at", "DATETIME", "8", "No", "Yes", "-", "Timestamp of message"]
        ]
    },
    {
        "title": "Table 16: SystemConfig (Global System Parameters)",
        "desc": "Stores system parameters such as weight limits and allowance rules.",
        "rows": [
            ["#", "Field name", "Type", "Size", "Unique", "Not Null", "PK/FK", "Notes"],
            ["1", "config_key", "VARCHAR", "50", "Yes", "Yes", "PK", "Unique config key (e.g., 'MAX_TOP_WEIGHT')"],
            ["2", "config_value", "VARCHAR", "255", "No", "Yes", "-", "Parameter value"],
            ["3", "description", "NVARCHAR", "255", "No", "No", "-", "Parameter description"],
            ["4", "updated_at", "DATETIME", "8", "No", "No", "-", "Last update timestamp"]
        ]
    },
    {
        "title": "Table 17: HorseRetirementRequest (Horse Retirement Management)",
        "desc": "Tracks formal retirement requests submitted by horse owners and processed by admin.",
        "rows": [
            ["#", "Field name", "Type", "Size", "Unique", "Not Null", "PK/FK", "Notes"],
            ["1", "id", "INT", "4", "Yes", "Yes", "PK", "Identity (1,1) auto-increment"],
            ["2", "horse_id", "INT", "4", "No", "Yes", "FK", "References Horse(id)"],
            ["3", "owner_id", "INT", "4", "No", "Yes", "FK", "References User(id)"],
            ["4", "reason", "NVARCHAR", "MAX", "No", "Yes", "-", "Detailed retirement reason"],
            ["5", "status", "VARCHAR", "30", "No", "Yes", "-", "Status ('PENDING', 'APPROVED', 'REJECTED')"],
            ["6", "admin_remarks", "NVARCHAR", "MAX", "No", "No", "-", "Admin decision notes"],
            ["7", "created_at", "DATETIME", "8", "No", "Yes", "-", "Request creation date"],
            ["8", "processed_at", "DATETIME", "8", "No", "No", "-", "Admin resolution date"]
        ]
    }
]

# Helper to build docx table matching template style
def add_custom_table(doc, tbl_info):
    p_title = doc.add_paragraph(tbl_info["title"])
    p_title.style = 'Heading 3'
    p_desc = doc.add_paragraph(tbl_info["desc"])
    
    rows_data = tbl_info["rows"]
    tbl = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    tbl.style = '_Style 16'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, r_data in enumerate(rows_data):
        row = tbl.rows[r_idx]
        for c_idx, val in enumerate(r_data):
            cell = row.cells[c_idx]
            cell.text = val
            # Header styling
            if r_idx == 0:
                p = cell.paragraphs[0]
                for run in p.runs:
                    run.font.bold = True
                    
    doc.add_paragraph("") # spacing

for tbl_info in missing_tables:
    add_custom_table(doc, tbl_info)

print("Added all 13 missing tables.")

# Save modified document
doc.save(doc_path)
print("Saved modified Template2_SRD Document.docx successfully!")
