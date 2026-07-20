import os
import sys
import docx

doc_path = r"folder_template/Template2_SRD Document.docx"
if not os.path.exists(doc_path):
    print("File not found:", doc_path)
    sys.exit(1)

doc = docx.Document(doc_path)
print("Loaded doc. Total paragraphs:", len(doc.paragraphs))

# Define the complete, cleanly structured Section II data
sec2_data = [
    {
        "subsystem": "1. Horse Owner Management",
        "fes": [
            {
                "code": "FE-01: Account Registration & Login",
                "trigger": "Navigation to the landing page and clicking 'Register' or 'Login'.",
                "desc": "Allows a Horse Owner to register their profile (email, password, fullname) and log in to manage their stable.",
                "logic": "Email and username must be unique and valid. Passwords must be hashed using BCrypt. Supports optional 2FA OTP verification via email."
            },
            {
                "code": "FE-02: Manage Owned Horses",
                "trigger": "Owner Dashboard -> 'My Stable' -> 'Register Horse' or 'Edit'.",
                "desc": "Add new horses, update profiles (breed, sex, date of birth, avatar), and track active status.",
                "logic": "Horse name must be unique. Date of birth must be in the past. Standard initial performance rating set to 52."
            },
            {
                "code": "FE-03: View Horse Profiles & Statistics",
                "trigger": "Owner Dashboard -> Clicking a horse card in the stable list.",
                "desc": "View complete horse stats: age, breed, current rating, total career starts, and win statistics.",
                "logic": "Displays dynamically calculated win rate: (total_wins / total_races) * 100."
            },
            {
                "code": "FE-04: View Seasons, Meetings & Schedules",
                "trigger": "Owner Dashboard -> 'Fixtures & Results'.",
                "desc": "View active season details, upcoming race meetings, and individual race schedules.",
                "logic": "Read-only access. Shows races that are SCHEDULED, DECLARATION_OPEN, or RUNNING."
            },
            {
                "code": "FE-05: Search Jockeys & Send Race Invitations",
                "trigger": "Selecting an upcoming race -> Clicking 'Invite Jockey' -> Selecting from active Jockey list.",
                "desc": "Search for eligible active jockeys and invite them to ride a specific horse in an upcoming race.",
                "logic": "Both horse and jockey must be registered for the corresponding Race Meeting. Horse rating must match the race rating class."
            },
            {
                "code": "FE-06: Manage Invitation Responses & Confirm Jockey",
                "trigger": "Owner Dashboard -> 'Invitations Sent'.",
                "desc": "View invitation status (PENDING, ACCEPTED, REJECTED, EXPIRED) and confirm the jockey.",
                "logic": "If jockey rejects, owner can re-invite. If accepted, status transitions to ENTRY: PENDING_ADMIN for Admin approval."
            },
            {
                "code": "FE-07: Confirm Horse Participation in Race Entry",
                "trigger": "Admin approves the entry, owner verifies the final gate assignment.",
                "desc": "Owner reviews the confirmed RaceEntry (jockey, horse, gate number).",
                "logic": "Weight allowance checked (e.g. sex allowance of 1.5kg for fillies/mares)."
            },
            {
                "code": "FE-08: Track Race Results & Earnings",
                "trigger": "Owner Dashboard -> 'History & Earnings'.",
                "desc": "View performance records, finishes, prize money earned, and rating adjustments after race publication.",
                "logic": "Prize money is calculated based on final position and race purse config."
            },
            {
                "code": "FE-09: Register Horses for Race Meetings",
                "trigger": "Owner Dashboard -> 'Race Meeting Registration'.",
                "desc": "Horse Owner registers eligible horses for an upcoming scheduled Race Meeting.",
                "logic": "Verifies horse status (must be ACTIVE) and ensures horse is not already registered for the same meeting. Creates a record in HorseRaceMeetingRegistration with PENDING status."
            },
            {
                "code": "FE-10: Submit Horse Retirement Request",
                "trigger": "Owner Dashboard -> My Horses -> 'Request Retirement'.",
                "desc": "Horse Owner submits a formal request with a detailed reason to retire an active horse from competition.",
                "logic": "Creates a record in HorseRetirementRequest with status PENDING. Horse cannot be entered into new races while retirement request is pending."
            }
        ]
    },
    {
        "subsystem": "2. Jockey Participation Management",
        "fes": [
            {
                "code": "FE-01: Account Registration & Login",
                "trigger": "Jockey registration form on landing page.",
                "desc": "Allows jockeys to register with their name, license info, and weight (kg) and log in.",
                "logic": "Jockey weight is critical for handicap weight calculation (must be between 52kg and 60kg)."
            },
            {
                "code": "FE-02: View Received Race Invitations",
                "trigger": "Jockey Dashboard -> 'New Invitations'.",
                "desc": "Display a list of pending invitations from horse owners, showing the race details, horse name, rating, and distance.",
                "logic": "Shows only invitations for races where the registration window is still open."
            },
            {
                "code": "FE-03: Accept or Reject Invitations",
                "trigger": "Clicking 'Accept' or 'Reject' buttons next to the invitation card.",
                "desc": "Jockey accepts to ride the horse, or rejects the invite.",
                "logic": "Accepting automatically creates a RaceEntry with status PENDING_ADMIN. Rejecting sets invite status to REJECTED."
            },
            {
                "code": "FE-04: View Race & Horse Details",
                "trigger": "Clicking on a confirmed race entry.",
                "desc": "View detailed profile of the assigned horse, the track characteristics (turf/dirt), and distance.",
                "logic": "Integrates with horse statistics API."
            },
            {
                "code": "FE-05: View Personal Race Schedule",
                "trigger": "Jockey Dashboard -> 'My Calendar'.",
                "desc": "View a calendar list of upcoming approved races they are scheduled to ride in.",
                "logic": "Jockeys cannot ride more than one horse per race."
            },
            {
                "code": "FE-06: View Personal Race Results & Achievements",
                "trigger": "Jockey Dashboard -> 'Performance History'.",
                "desc": "Track total races ridden, number of top-3 finishes, win rate, and total points earned across seasons.",
                "logic": "Updates total_races_participated and total_top3_finishes upon official race result publication."
            },
            {
                "code": "FE-07: Register Availability for Race Meetings",
                "trigger": "Jockey Dashboard -> 'Race Meeting Registration'.",
                "desc": "Jockey registers availability to ride in races scheduled for an upcoming Race Meeting.",
                "logic": "Verifies jockey account status is ACTIVE. Creates a record in JockeyRaceMeetingRegistration with PENDING status for Admin review."
            }
        ]
    },
    {
        "subsystem": "3. Race Officiating System",
        "fes": [
            {
                "code": "FE-01: Verify Horse Information & Weight",
                "trigger": "Referee Dashboard -> 'Pre-Race Weigh-In'.",
                "desc": "Referee verifies the jockey's actual weight and the horse's identity prior to post position entry.",
                "logic": "Calculates carried weight (Jockey weight + handicap weight). Verifies minimum weight requirements."
            },
            {
                "code": "FE-02: Monitor Race Progress",
                "trigger": "Referee Dashboard -> Clicking 'Start Race' -> Livestreaming interface.",
                "desc": "Provides the referee with controls to monitor the live race stream and manage race states.",
                "logic": "Transitions race status from RACE_ASSIGNED to RUNNING."
            },
            {
                "code": "FE-03: Record & Handle Rule Violations",
                "trigger": "Clicking 'Report Incident' during or immediately after the race.",
                "desc": "Referee records rules infractions (e.g. cutting lane, whip abuse) and penalties (disqualification, fines).",
                "logic": "A serious violation transitions RaceEntry status to DISQUALIFIED and results in final_position being null/disqualified."
            },
            {
                "code": "FE-04: Approve & Confirm Official Race Results",
                "trigger": "Referee Dashboard -> 'Result Entry'.",
                "desc": "Referee enters the final position and finish times for all horses, confirms the result.",
                "logic": "Transitions race status to OFFICIAL and triggers automatic rating updates and prize money calculation in the database."
            },
            {
                "code": "FE-05: Create & Store Race Reports",
                "trigger": "After race results are published, clicking 'Create Report'.",
                "desc": "Form to save the official steward's report, summarizing violations, weather conditions, track status, and general comments.",
                "logic": "Saves to steward_report column in Race table."
            }
        ]
    },
    {
        "subsystem": "4. Spectator Management System",
        "fes": [
            {
                "code": "FE-01: Register & Log In",
                "trigger": "Spectator signup button on the main page.",
                "desc": "Registration of spectators for tracking favorite horses/jockeys and accessing AI chatbot history.",
                "logic": "Simple email registration, does not require weight validation."
            },
            {
                "code": "FE-02: View Season, Meeting & Race Information",
                "trigger": "Portal Landing -> 'Schedules & Info'.",
                "desc": "Spectators browse current season standings, race meetings, list of races, and historical results.",
                "logic": "Read-only access to published data."
            },
            {
                "code": "FE-03: Track Live Results & Live Chat",
                "trigger": "Portal Landing -> 'Watch Live'.",
                "desc": "Watch live video streams of active races, view real-time gate positions, and participate in the public chat.",
                "logic": "Connects to backend via WebSocket (/ws/chat/{raceId})."
            },
            {
                "code": "FE-04: AI Race Outcome Predictor & Smart Assistant",
                "trigger": "Landing Portal -> 'AI Race Assistant / Predictor'.",
                "desc": "Spectator interacts with the AI Assistant to query race stats and obtain AI-driven predictions for upcoming race outcomes.",
                "logic": "Integrates with Python backend AI engine (chatbot.py, predictor.py) utilizing multi-key API rotation (Gemini & Groq) for real-time statistical predictions based on horse ratings, jockey win rates, and track conditions."
            }
        ]
    },
    {
        "subsystem": "5. Race Management & Administration System",
        "fes": [
            {
                "code": "FE-01: Manage User Accounts",
                "trigger": "Admin Dashboard -> 'User Management'.",
                "desc": "List, edit, deactivate, or activate users (Owners, Jockeys, Referees, Spectators).",
                "logic": "Admins cannot delete their own account."
            },
            {
                "code": "FE-02: Assign Role Permissions",
                "trigger": "Admin Dashboard -> 'System Configuration'.",
                "desc": "Configure and review security permissions and configuration values (e.g. MAX_TOP_WEIGHT, WEIGHT_PER_POINT).",
                "logic": "Persists settings in SystemConfig table."
            },
            {
                "code": "FE-03: Manage Seasons, Meetings & Races",
                "trigger": "Admin Dashboard -> 'Tournament Management'.",
                "desc": "Create and edit Seasons, add Race Meetings, and schedule individual races.",
                "logic": "Start date must be before end date. A race meeting must belong to an active season."
            },
            {
                "code": "FE-04: Manage Season Class Rules",
                "trigger": "Admin Dashboard -> 'Class Rules'.",
                "desc": "Define rating brackets (min/max rating) and purse values for each race class (Class 1 to Class 5).",
                "logic": "Rules must not overlap in rating brackets."
            },
            {
                "code": "FE-05: Approve Owner & Jockey Registrations",
                "trigger": "Admin Dashboard -> 'Registrations' -> 'Approve / Reject'.",
                "desc": "Review and approve jockeys and owners registering for a specific race meeting.",
                "logic": "Checks eligibility and licensing status."
            },
            {
                "code": "FE-06: Assign Referees to Races",
                "trigger": "Admin Dashboard -> 'Assign Referees'.",
                "desc": "Select and assign certified referees to oversee scheduled races.",
                "logic": "A race can have one or more referees."
            },
            {
                "code": "FE-07: Approve Race Entries",
                "trigger": "Admin Dashboard -> 'Race Card Approval'.",
                "desc": "Approve the final entries of jockeys and horses for a race after the declaration deadline closes.",
                "logic": "Automatically generates gate numbers randomly (1 to 14) and transitions entry status to APPROVED."
            },
            {
                "code": "FE-08: Publish Official Race Results",
                "trigger": "Admin Dashboard -> 'Publish Results'.",
                "desc": "Reviews results entered by the referee and publishes them globally.",
                "logic": "Triggers season rating updates and unlocks prize earnings."
            },
            {
                "code": "FE-09: Manage Race Meeting Registrations",
                "trigger": "Admin Dashboard -> 'Meeting Registrations' -> Approve / Reject.",
                "desc": "Admin reviews and approves or rejects registrations from Horse Owners, Jockeys, and Horses for upcoming Race Meetings.",
                "logic": "Updates registration status in OwnerRaceMeetingRegistration, JockeyRaceMeetingRegistration, and HorseRaceMeetingRegistration. Only approved participants can be invited or declared for races in that meeting."
            },
            {
                "code": "FE-10: Process Horse Retirement Requests",
                "trigger": "Admin Dashboard -> 'Retirement Requests'.",
                "desc": "Admin reviews pending horse retirement requests submitted by owners, adds administrative remarks, and approves or rejects them.",
                "logic": "Approving sets HorseRetirementRequest status to APPROVED and updates Horse status to RETIRED. Rejection returns horse status to ACTIVE."
            },
            {
                "code": "FE-11: Configure Live Stream Settings",
                "trigger": "Admin Dashboard -> 'Live Settings'.",
                "desc": "Admin configures live streaming options and sets the YouTube Live URL for scheduled races.",
                "logic": "Validates URL format and updates youtube_live_url in Race table. Embedded player on spectator dashboard updates automatically."
            }
        ]
    }
]

# Find start and end paragraph of Section II
sec2_start_idx = None
sec3_start_idx = None

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == "II. Functional Requirements":
        sec2_start_idx = i
    elif t == "III. Database Design":
        sec3_start_idx = i

print(f"Section II starts at P{sec2_start_idx}, Section III starts at P{sec3_start_idx}")

if sec2_start_idx is None or sec3_start_idx is None:
    print("Could not find section boundaries!")
    sys.exit(1)

# Target paragraph to insert BEFORE (the III. Database Design paragraph)
target_p = doc.paragraphs[sec3_start_idx]

# Collect all paragraphs between sec2_start_idx + 1 and sec3_start_idx - 1 to remove
to_remove = doc.paragraphs[sec2_start_idx + 1 : sec3_start_idx]

print(f"Removing {len(to_remove)} old paragraphs from Section II...")
for p in to_remove:
    p._element.getparent().remove(p._element)

print("Old Section II content removed.")

# Now insert new Section II paragraphs BEFORE target_p in exact forward order
def insert_before(target_paragraph, text, style_name='Normal'):
    new_p = target_paragraph.insert_paragraph_before(text, style=style_name)
    return new_p

for sys_item in sec2_data:
    # Subsystem Heading 2
    insert_before(target_p, sys_item["subsystem"], style_name='Heading 2')
    
    for fe in sys_item["fes"]:
        # Heading 3
        insert_before(target_p, fe["code"], style_name='Heading 3')
        # Trigger
        insert_before(target_p, f"• Function trigger: {fe['trigger']}", style_name='Normal')
        # Description
        insert_before(target_p, f"• Function description: {fe['desc']}", style_name='Normal')
        # Business logic
        insert_before(target_p, f"• Business logic & validations: {fe['logic']}", style_name='Normal')
        # Spacer
        insert_before(target_p, "", style_name='Normal')

doc.save(doc_path)
print("Successfully saved updated Template2_SRD Document.docx!")
