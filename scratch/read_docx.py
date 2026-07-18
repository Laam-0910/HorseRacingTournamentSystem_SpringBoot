import os
import sys
import subprocess

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

doc_path = r"d:\FPTU\k5\4swp\HorseRacingTournamentSystem_SpringBoot\folder_template\Template2_SRD Document.docx"

if not os.path.exists(doc_path):
    print(f"File not found at: {doc_path}")
    sys.exit(1)

print("Reading document...")
doc = docx.Document(doc_path)

# Let's search for the section heading "4. System Functions" or similar
found_section = False
section_paragraphs = []

print("Document loaded. Total paragraphs:", len(doc.paragraphs))

# Print headings to see structure
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    # Look for headings starting with "4." or containing "System Functions"
    if p.style.name.startswith("Heading") or (text.lower().startswith("4.") or "system functions" in text.lower()):
        print(f"[{p.style.name}] {text}")

# Let's also print some text around 4. System Functions to see if there's any template instructions
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "4." in text and "System Functions" in text:
        print(f"\n--- FOUND SECTION 4 AT PARAGRAPH {i} ---")
        # Print next 20 paragraphs
        for j in range(i, min(i + 40, len(doc.paragraphs))):
            p_text = doc.paragraphs[j].text.strip()
            if p_text:
                print(f"{j}: [{doc.paragraphs[j].style.name}] {p_text}")
        break
