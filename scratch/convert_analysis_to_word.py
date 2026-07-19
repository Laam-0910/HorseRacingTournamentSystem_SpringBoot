import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

doc_path = r"System_Architecture_and_Feature_Analysis.docx"
doc = docx.Document()

# Page setup: Standard margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Style helpers
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121) # #1F4E79
    return p

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(89, 89, 89)
    return p

def add_heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    return p

def add_heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(46, 117, 182) # #2E75B6
    return p

def add_heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 51, 51)
    return p

def add_callout(text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "EBF3FA")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border thick blue
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:left w:val="single" w:sz="36" w:space="0" w:color="2E75B6"/>
            <w:top w:val="none"/>
            <w:right w:val="none"/>
            <w:bottom w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    doc.add_paragraph("") # spacing

# Document Header
add_title("BÁO CÁO PHÂN TÍCH KIẾN TRÚC & CODEBASE 17 CHỨC NĂNG")
add_subtitle("HORSE RACING TOURNAMENT MANAGEMENT SYSTEM (SWP391 - GROUP 4)")

add_callout("Tài liệu này tổng hợp và phân tích chi tiết vị trí mã nguồn (Frontend React, Backend Spring Boot, Database SQL Server, Python AI Engine), luồng xử lý dữ liệu và cơ chế kỹ thuật chuyên sâu cho toàn bộ 17 chức năng trong Project Tracking.")

# SECTION 1
add_heading1("I. TỔNG QUAN KIẾN TRÚC HỆ THỐNG (SYSTEM ARCHITECTURE OVERVIEW)")

p_arch = doc.add_paragraph("Hệ thống được thiết kế theo kiến trúc Phân tầng (Layered Micro-services Hybrid) kết hợp giữa Monolith Backend (Spring Boot), Single Page Application (React + Vite) và Python AI Microservice:")
p_arch.paragraph_format.space_after = Pt(6)

arch_items = [
    ("1. Frontend Layer (Port 5173): ", "React 18 + Vite + TypeScript. Quản lý trạng thái bằng AuthContext, giao diện Glassmorphism Modern UI và đa ngôn ngữ i18n."),
    ("2. Backend Layer (Port 8080): ", "Java 17 + Spring Boot. Cung cấp hệ thống RESTful APIs, Spring Security Stateless JWT Filter, WebSocket Server (/ws/chat) và Scheduler tự động."),
    ("3. Python AI Engine (Port 5000): ", "Python 3.10 + Flask. Độc lập xử lý AI Chatbot RAG (kết hợp xoay vòng Gemini/Groq API Keys) và Mô hình Machine Learning Predictor."),
    ("4. Database Layer: ", "Microsoft SQL Server (HorseRacingDB) với 17 bảng được thiết kế chuẩn hóa 3NF.")
]

for title, desc in arch_items:
    p = doc.add_paragraph()
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.color.rgb = RGBColor(46, 117, 182)
    r2 = p.add_run(desc)

# SECTION 2: MATRIX TABLE
add_heading1("II. BẢN ĐỒ MA TRẬN 17 CHỨC NĂNG (FEATURE MAPPING MATRIX)")

matrix_headers = ["#", "Tên Chức Năng", "Vai Trò", "Component Frontend", "Backend Service/Controller", "Bảng Database SQL"]
matrix_data = [
    ["1", "User Login & Authentication", "All Users", "Login.tsx, VerifyLogin.tsx", "AuthController, AuthService, JwtTokenProvider", "User, Role"],
    ["2", "Manage Users", "Admin", "Users.tsx, UserEdit.tsx", "AdminUserController, AdminUserService", "User, Role"],
    ["3", "Horse Management (CRUD)", "Admin, Owner", "Horses.tsx, HorseOwner.tsx", "HorseOwnerController, HorseService", "Horse, User"],
    ["4", "Horse Retirement", "Admin, Owner", "AdminHorseRetirement.tsx", "HorseRetirementController, HorseRetirementService", "HorseRetirementRequest, Horse"],
    ["5", "Jockey Management", "Admin, Jockey", "Jockey.tsx", "JockeyController, JockeyService", "User, JockeyRaceMeetingRegistration"],
    ["6", "Manage Race & Tournaments", "Admin", "Season.tsx, RaceMeeting.tsx, Race.tsx", "SeasonController, RaceMeetingController, RaceController", "Season, SeasonClassRule, RaceMeeting, Race"],
    ["7", "View Public Schedules & Results", "Guest, All", "Landing.tsx, Fixtures.tsx, Results.tsx", "PublicController, RaceDayScheduleService", "RaceMeeting, Race, RaceEntry"],
    ["8", "Referee Actions", "Referee", "RefereeHub.tsx, RefereeDuties.tsx", "RefereeController, RefereeService", "Race, RaceEntry, Violation, RaceReferee"],
    ["9", "Public Chat", "All Users", "Chatbot.tsx", "PublicChatController, ChatMessageRepository", "ChatMessage"],
    ["10", "AI Chatbot Integration & Predictor", "Guest, All", "Chatbot.tsx", "app.py, chatbot.py, predictor.py (Python)", "Horse, User, RaceEntry"],
    ["11", "Spectator Dashboard Charts", "Spectator", "Spectator.tsx, Statistics.tsx", "PublicController, DashboardUtils", "RaceEntry, Horse, User"],
    ["12", "Real-time Livestream Chat", "Spectator", "Livestream.tsx, ViewLive.tsx", "ChatWebSocketHandler, WebSocketConfig", "ChatMessage, Race"],
    ["13", "Server-side Pagination & Search", "Admin, All", "All Workflow Views", "Custom JPQL / Pageable Repositories", "All 17 Tables"],
    ["14", "Stewards' Inquiry & Weight Check", "Referee, Admin", "RefereeCheck.tsx, IncidentReport.tsx", "RefereeService, ProcessResultsService", "RaceEntry, SystemConfig, Violation"],
    ["15", "Automated Race Cancellation", "System / Admin", "Backend Scheduler Task", "RaceStatusScheduler, RaceService", "Race, RaceInvitation, RaceEntry"],
    ["16", "Multilingual System (i18n)", "All Users", "i18n.ts, dictionary.json", "SystemConfigController", "SystemConfig"],
    ["17", "UI/UX Framework & Design System", "All Users", "index.css, DashboardLayout.tsx", "N/A (Frontend Framework)", "N/A"]
]

tbl_m = doc.add_table(rows=len(matrix_data) + 1, cols=6)
tbl_m.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr_cells = tbl_m.rows[0].cells
for idx, text in enumerate(matrix_headers):
    hdr_cells[idx].text = text
    set_cell_background(hdr_cells[idx], "1F4E79") # Deep Blue
    set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=100, right=100)
    p = hdr_cells[idx].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)

# Data rows
for r_idx, row_data in enumerate(matrix_data):
    row_cells = tbl_m.rows[r_idx + 1].cells
    bg_color = "F2F4F8" if r_idx % 2 == 1 else "FFFFFF"
    for c_idx, val in enumerate(row_data):
        cell = row_cells[c_idx]
        cell.text = val
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)

doc.add_paragraph("") # Spacing

# SECTION 3: DETAILED 17 FEATURES ANALYSIS
add_heading1("III. PHÂN TÍCH CHI TIẾT 17 CHỨC NĂNG (DETAILED FEATURE BREAKDOWN)")

features_detail = [
    {
        "title": "1. User Login & Authentication (Đăng nhập & Xác thực 2 Lớp)",
        "frontend": "frontend/src/app/components/auth/Login.tsx, VerifyLogin.tsx, AuthContext.tsx",
        "backend": "com.horseracing.backend.controller.AuthController, AuthService.java, JwtTokenProvider.java",
        "db": "Bảng [User], Role",
        "desc": "Hệ thống xác thực người dùng bằng username/password qua BCrypt. Nếu tài khoản cấu hình require_otp = 1, hệ thống sinh mã OTP 6 số ngẫu nhiên gửi qua EmailSender. Sau khi xác thực thành công, JwtTokenProvider khởi tạo JWT Token 24h chứa UserID và Role. Token được nạp vào SecurityContextHolder qua JwtAuthenticationFilter."
    },
    {
        "title": "2. Manage Users (Quản lý Người dùng - Admin)",
        "frontend": "frontend/src/app/components/admin-workflow/Users.tsx, UserEdit.tsx",
        "backend": "AdminUserController.java, AdminUserService.java, UserRepository.java",
        "db": "Bảng [User], Role",
        "desc": "Cho phép Admin tìm kiếm, lọc người dùng theo Role và trạng thái ACTIVE/INACTIVE. Chỉnh sửa thông tin tài khoản, cân nặng nài ngựa và vô hiệu hóa tài khoản. Quy tắc bảo mật: AdminUserService kiểm tra ngăn cản Admin tự khóa tài khoản của chính mình."
    },
    {
        "title": "3. Horse Management - CRUD (Quản lý Chiến mã)",
        "frontend": "Horses.tsx, HorseOwner.tsx",
        "backend": "HorseOwnerController.java, HorseService.java, HorseRepository.java",
        "db": "Bảng Horse",
        "desc": "Quản lý hồ sơ ngựa đua: Tên, giống, giới tính, ngày sinh, điểm Rating khởi tạo (52), ảnh đại diện Base64. Khi Owner khởi tạo ngựa, backend tự động lấy OwnerID từ JWT Token để đảm bảo ngựa thuộc quyền sở hữu duy nhất của chủ sở hữu đó."
    },
    {
        "title": "4. Horse Retirement (Giải nghệ Ngựa)",
        "frontend": "AdminHorseRetirement.tsx",
        "backend": "HorseRetirementController.java, HorseRetirementService.java",
        "db": "Bảng HorseRetirementRequest, Horse",
        "desc": "Chủ ngựa gửi đơn xin giải nghệ ngựa kèm lý do. Hệ thống tạo bản ghi PENDING trong HorseRetirementRequest. Admin xem xét, ghi nhận xét (admin_remarks) và bấm Approve/Reject. Nếu Approve, trạng thái ngựa tự động chuyển thành RETIRED."
    },
    {
        "title": "5. Jockey Management (Quản lý Nài Ngựa)",
        "frontend": "Jockey.tsx",
        "backend": "JockeyController.java, JockeyService.java",
        "db": "Bảng [User], JockeyRaceMeetingRegistration, RaceInvitation",
        "desc": "Jockey cập nhật cân nặng cá nhân (phục vụ tính Tạ gánh Carried Weight) và đăng ký tham gia Ngày đua (Race Meeting). Đăng ký được duyệt thì Owner mới có thể gửi lời mời thi đấu. Bấm Accept lời mời sẽ tạo bản ghi RaceEntry PENDING_ADMIN."
    },
    {
        "title": "6. Manage Race & Tournaments (Quản lý Giải đấu & Trận đua)",
        "frontend": "Season.tsx, RaceMeeting.tsx, Race.tsx",
        "backend": "SeasonController.java, RaceMeetingController.java, RaceController.java",
        "db": "Bảng Season, SeasonClassRule, RaceMeeting, Race",
        "desc": "Admin quản lý Mùa giải (Season) và quy định 5 Hạng đua Rating Class Rules (Class 1 -> Class 5). Thiết lập thông số trận đua: cự ly (1000m-2400m), mặt sân (Turf/Dirt), thời gian mở/chốt sổ đăng ký và giới hạn số ngựa (min 3, max 14)."
    },
    {
        "title": "7. View Public Schedules & Results (Xem Lịch đua & Kết quả Công khai)",
        "frontend": "Landing.tsx, Fixtures.tsx, Results.tsx",
        "backend": "PublicController.java, RaceDayScheduleService.java",
        "db": "Bảng RaceMeeting, Race, RaceEntry, Horse",
        "desc": "Khán giả và khách xem (Guest) không cần đăng nhập có thể tra cứu toàn bộ lịch đua sắp diễn ra, danh sách ngựa tham gia (Racecard) và thứ hạng/thời gian về đích của các trận đua đã hoàn thành kèm đồng hồ đếm ngược Live Countdown."
    },
    {
        "title": "8. Referee Actions (Chức năng Trọng tài Đua)",
        "frontend": "RefereeHub.tsx, RefereeDuties.tsx, RefereeCheck.tsx",
        "backend": "RefereeController.java, RefereeService.java",
        "db": "Bảng Race, RaceEntry, Violation, RaceReferee",
        "desc": "Trọng tài cân đo trước trận (Pre-race weigh-in), giám sát trận đua trực tiếp, ghi nhận các lỗi vi phạm (Violation) như chèn làn, phạt roi. Sau trận đua, Trọng tài nhập kết quả thứ hạng, thời gian về đích và viết báo cáo giám sát steward_report."
    },
    {
        "title": "9. Public Chat (Kênh Trò chuyện Cộng đồng)",
        "frontend": "Chatbot.tsx",
        "backend": "PublicChatController.java, ChatMessageRepository.java",
        "db": "Bảng ChatMessage",
        "desc": "Kênh trò chuyện cộng đồng cho khán giả trao đổi thông tin xung quanh các trận đua. Tin nhắn được lưu lịch sử kèm mốc thời gian sent_at vào bảng ChatMessage."
    },
    {
        "title": "10. AI Chatbot Integration & Predictor (Trợ lý AI & Dự đoán Trận đua)",
        "frontend": "Chatbot.tsx",
        "backend": "app.py, chatbot.py, predictor.py (Python Flask Engine)",
        "db": "Bảng Horse, User, RaceEntry (pyodbc)",
        "desc": "Tích hợp Python AI Microservice độc lập. Chatbot RAG hỗ trợ hỏi đáp với cơ chế xoay vòng danh sách GEMINI_API_KEYS và GROQ_API_KEYS chống chạm trần Rate Limit. Mô hình Gradient Boosting (predictor.py) dự đoán xác suất % lọt Top 3 của chiến mã."
    },
    {
        "title": "11. Spectator Dashboard Charts (Biểu đồ Thống kê Khán giả)",
        "frontend": "Spectator.tsx, Statistics.tsx",
        "backend": "PublicController.java, DashboardUtils.java",
        "db": "Bảng RaceEntry, Horse, User",
        "desc": "Trực quan hóa dữ liệu thống kê phong độ đua ngựa dưới dạng SVG/CSS Charts responsive: Tỷ lệ thắng theo dòng ngựa (Breed win rate), biểu đồ biến động điểm Rating qua các giải đua, xếp hạng Top Nài ngựa."
    },
    {
        "title": "12. Real-time Livestream Chat (Chat Trực tiếp theo thời gian thực)",
        "frontend": "Livestream.tsx, ViewLive.tsx",
        "backend": "ChatWebSocketHandler.java, WebSocketConfig.java",
        "db": "Bảng ChatMessage, Race",
        "desc": "Khán giả theo dõi trận đua live qua video YouTube nhúng và tham gia phòng chat trực tiếp. Sử dụng kết nối Full-Duplex WebSocket (/ws/chat/{raceId}) giúp tin nhắn được phát tức thì tới tất cả người dùng trong phòng đua."
    },
    {
        "title": "13. Server-side Pagination & Searching (Phân trang & Tìm kiếm Server)",
        "frontend": "Tất cả các màn hình danh sách Admin / Owner / Referee",
        "backend": "Custom JPQL / Spring Data Pageable Repositories",
        "db": "Tất cả 17 bảng Database",
        "desc": "Mọi truy vấn dữ liệu lớn (danh sách người dùng, ngựa, trận đua) đều sử dụng tham số phân trang page, size, sort, keyword ở mức SQL Server (OFFSET...FETCH), giúp tối ưu hóa bộ nhớ RAM server và giảm thời gian tải trang."
    },
    {
        "title": "14. Stewards' Inquiry & Weight Comparison (Kiểm tra Tạ gánh Trọng tài)",
        "frontend": "RefereeCheck.tsx, IncidentReport.tsx",
        "backend": "RefereeService.java, ProcessResultsService.java",
        "db": "Bảng SystemConfig, RaceEntry, Violation",
        "desc": "Tính toán Tạ gánh tự động: CarriedWeight = JockeyWeight + HandicapWeight - SexAllowance. Nếu chênh lệch cân nặng thực tế sau trận (Weigh-out) vượt quá MAX_OVERWEIGHT_ALLOWED (1.0kg), hệ thống phát cảnh báo Stewards' Inquiry cờ đỏ."
    },
    {
        "title": "15. Automated Race Cancellation (Tự động Hủy Trận đua)",
        "frontend": "Giao diện quản lý lịch đua",
        "backend": "RaceStatusScheduler.java, RaceService.java",
        "db": "Bảng Race, RaceInvitation, RaceEntry",
        "desc": "Timer task định kỳ RaceStatusScheduler kiểm tra các trận đua sau khi chốt sổ đăng ký. Nếu số lượng ngựa được duyệt < min_entries (dưới 3 ngựa), trận đua tự động chuyển sang CANCELLED và giải phóng lịch cho nài/chủ ngựa."
    },
    {
        "title": "16. Multilingual System - i18n (Hệ thống Đa ngôn ngữ)",
        "frontend": "i18n.ts, dictionary.json",
        "backend": "SystemConfigController.java",
        "db": "Bảng SystemConfig",
        "desc": "Cung cấp bộ từ điển dịch đa ngôn ngữ linh hoạt (Tiếng Anh, Tiếng Việt) cho toàn bộ nhãn giao diện, thông báo lỗi, menu và tên trạng thái mà không cần reload trang."
    },
    {
        "title": "17. UI/UX Framework & Design System (Hệ thống Giao diện Modern)",
        "frontend": "index.css, DashboardLayout.tsx, confirm.tsx",
        "backend": "N/A (Frontend Framework)",
        "db": "N/A",
        "desc": "Phong cách thiết kế Glassmorphic Modern UI sang trọng, phối màu Dark Slate & Royal Gold, hiệu ứng mờ kính (Backdrop Blur), animation chuyển trang mượt mà và hộp thoại xác nhận confirm.tsx tùy biến."
    }
]

for item in features_detail:
    add_heading2(item["title"])
    
    p_loc = doc.add_paragraph()
    r_f = p_loc.add_run("• Frontend: ")
    r_f.bold = True
    r_f.font.color.rgb = RGBColor(46, 117, 182)
    p_loc.add_run(item["frontend"])
    
    p_b = doc.add_paragraph()
    r_b = p_b.add_run("• Backend: ")
    r_b.bold = True
    r_b.font.color.rgb = RGBColor(46, 117, 182)
    p_b.add_run(item["backend"])
    
    p_d = doc.add_paragraph()
    r_d = p_d.add_run("• Database: ")
    r_d.bold = True
    r_d.font.color.rgb = RGBColor(46, 117, 182)
    p_d.add_run(item["db"])
    
    p_desc = doc.add_paragraph()
    r_c = p_desc.add_run("• Môt tả & Logic: ")
    r_c.bold = True
    r_c.font.color.rgb = RGBColor(31, 78, 121)
    p_desc.add_run(item["desc"])

# Save Word Document
doc.save(doc_path)
print("Successfully generated System_Architecture_and_Feature_Analysis.docx!")
