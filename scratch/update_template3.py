import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc_path = r"folder_template/Template3_Final Release Document.docx"
if not os.path.exists(doc_path):
    print("File not found:", doc_path)
    sys.exit(1)

doc = docx.Document(doc_path)

# Clear existing paragraphs and tables
for p in list(doc.paragraphs):
    p._element.getparent().remove(p._element)

for tbl in list(doc.tables):
    tbl._element.getparent().remove(tbl._element)

# Helper to add styled title
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    return p

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(89, 89, 89)
    return p

def add_heading1(text):
    p = doc.add_paragraph(style='Heading 1')
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    return p

def add_heading2(text):
    p = doc.add_paragraph(style='Heading 2')
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(46, 117, 182)
    return p

def add_heading3(text):
    p = doc.add_paragraph(style='Heading 3')
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    return p

# Cover Page / Header
add_title("Horse Racing Tournament System")
add_subtitle("Final Release Document")

p_meta = doc.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_meta.add_run("Course: SWP391 - Software Development Project | Class: SE1918 | Group: Group 4\nLecturer: Đỗ Tấn Nhàn | Date: July 2026\n– Hanoi, Vietnam –\n\n")
r.font.size = Pt(11)
r.font.italic = True

# Section I: Deliverable Package
add_heading1("I. Deliverable Package")

p_intro = doc.add_paragraph("This section outlines all source programs, configuration scripts, database artifacts, and documentation deliverables included in the final release package of the Horse Racing Tournament System.")

# Deliverables Table
table_data = [
    ["No.", "Deliverable Artifact", "File Name / Location", "Notes & Description"],
    ["1", "Database SQL Script", "SQLHorseRacing.sql", "Complete MS SQL Server database schema, identity constraints, check constraints, foreign keys, and seed data for initial setup."],
    ["2", "AI Usage Report", "Template0_AI Usage Report.xlsx", "Detailed log of AI tools (Gemini, Antigravity, ChatGPT) usage across SDLC phases, prompts, validation metrics, and link evidence."],
    ["3", "Project Tracking Matrix", "Template1_Project Tracking.xlsx", "Matrix tracking 17 core features/screens, team role assignments, development complexity, and final 100% completion status."],
    ["4", "System Requirement & Design", "Template2_SRD Document.docx", "Full SRS/SDS document containing system overview, ERD diagrams, 17 Database Table Details, and functional requirements for 5 user roles."],
    ["5", "Final Release Document", "Template3_Final Release Document.docx", "Release manifest, deliverable package tracking, system installation guide, and end-user manual."],
    ["6", "Spring Boot Backend Source", "/backend/", "Java 17 Spring Boot backend application handling REST APIs, JPA repositories, security authorization, and WebSocket handlers."],
    ["7", "React Frontend Source", "/frontend/", "React / Vite frontend web application implementing responsive dashboards for Admin, Owner, Jockey, Referee, and Spectator."],
    ["8", "Python AI Engine Source", "/backend/.../python/", "Python Flask microservice (app.py, chatbot.py, predictor.py) providing AI Chatbot Q&A and Machine Learning race outcome predictions."]
]

tbl = doc.add_table(rows=len(table_data), cols=4)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for r_idx, row_data in enumerate(table_data):
    row = tbl.rows[r_idx]
    for c_idx, val in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.text = val
        if r_idx == 0:
            for run in cell.paragraphs[0].runs:
                run.font.bold = True

doc.add_paragraph("")

p_rel = doc.add_paragraph()
p_rel.add_run("Repository & Demonstration Links:\n").bold = True
p_rel.add_run("• Source Code Repository: ").bold = True
p_rel.add_run("https://github.com/Laam-0910/HorseRacingTournamentSystem_SpringBoot\n")
p_rel.add_run("• Target Git Branch: ").bold = True
p_rel.add_run("main / Laam\n")
p_rel.add_run("• Live Technology Stack: ").bold = True
p_rel.add_run("Java 17, Spring Boot, React, Vite, MS SQL Server, Python 3.10, WebSocket, Tailwind CSS / Vanilla CSS.")

# Section II: Installation Guides
add_heading1("II. Installation & System Deployment Guide")

doc.add_paragraph("This guide provides step-by-step instructions to configure, build, and deploy the Horse Racing Tournament System across all backend, frontend, database, and AI microservice components.")

add_heading2("1. System Prerequisites")
p_req = doc.add_paragraph()
p_req.add_run("Ensure the following runtime environments are installed on the target machine:\n")
p_req.add_run("• Java Development Kit (JDK): Version 17 or higher\n")
p_req.add_run("• Build Tool: Apache Maven 3.8+ (or included ./mvnw wrapper)\n")
p_req.add_run("• Node.js & npm: Node v18.0.0+ and npm v9.0.0+\n")
p_req.add_run("• Python Runtime: Python 3.10+ (with pip package manager)\n")
p_req.add_run("• Database Engine: Microsoft SQL Server 2019+ (Express/Developer Edition) or Azure SQL\n")
p_req.add_run("• IDE / Editor: Visual Studio Code, IntelliJ IDEA, or Eclipse")

add_heading2("2. Database Initialization")
doc.add_paragraph("1. Open Microsoft SQL Server Management Studio (SSMS) or SQLCMD.\n"
                  "2. Open and execute the SQL script SQLHorseRacing.sql located in the project root directory.\n"
                  "3. The script creates the database HorseRacingDB, establishes 17 relational tables with constraints, and inserts initial seed data (default roles, system configs, sample users, seasons, and horses).")

add_heading2("3. Backend Configuration & Execution (Spring Boot)")
doc.add_paragraph("1. Navigate to the backend root directory:\n"
                  "   cd backend\n"
                  "2. Configure database connection credentials in application.properties or via environment variables (.env):\n"
                  "   spring.datasource.url=jdbc:sqlserver://localhost:1433;databaseName=HorseRacingDB;encrypt=false\n"
                  "   spring.datasource.username=sa\n"
                  "   spring.datasource.password=YourPassword\n"
                  "3. Build and launch the Spring Boot server:\n"
                  "   ./mvnw spring-boot:run\n"
                  "4. The backend service will start on http://localhost:8080.")

add_heading2("4. Python AI Service Setup (Chatbot & Predictor)")
doc.add_paragraph("1. Navigate to the Python microservice directory:\n"
                  "   cd backend/src/main/webapp/WEB-INF/python\n"
                  "2. Install required Python packages:\n"
                  "   pip install -r requirements.txt\n"
                  "3. Configure your API keys in the environment file (.env):\n"
                  "   GEMINI_API_KEYS=key1,key2,key3\n"
                  "   GROQ_API_KEYS=key1,key2\n"
                  "4. Start the Flask AI engine:\n"
                  "   python app.py\n"
                  "5. The AI microservice listens on http://127.0.0.1:5000.")

add_heading2("5. Frontend Application Deployment (React / Vite)")
doc.add_paragraph("1. Open a terminal and navigate to the frontend directory:\n"
                  "   cd frontend\n"
                  "2. Install JavaScript dependencies:\n"
                  "   npm install\n"
                  "3. Launch the development server:\n"
                  "   npm run dev\n"
                  "4. Open your browser and navigate to http://localhost:5173 to access the web application.")

# Section III: User Manual
add_heading1("III. User Manual & System Workflows")

doc.add_paragraph("The Horse Racing Tournament System caters to 5 distinct user roles. Below are detailed operational guides for each key workflow.")

add_heading2("1. System Overview & Role Authorization")
doc.add_paragraph("• Administrator: Full governance over seasons, class rules, user accounts, race schedules, referee assignments, registration approvals, and publishing official race results.\n"
                  "• Horse Owner: Manages horse stable profiles, submits horse retirement requests, registers horses for race meetings, and sends race invitations to certified jockeys.\n"
                  "• Jockey: Updates weight profile, registers race meeting availability, accepts or declines race invitations from owners, and views assigned starting gates.\n"
                  "• Race Referee: Conducts pre-race weigh-in verification, oversees live race execution, records rule infractions/penalties, and submits official finish positions.\n"
                  "• Spectator: Browses public schedules, tracks live race video streams, participates in real-time WebSocket chat, and queries the AI Chatbot Predictor.")

add_heading2("2. Workflow 1: Tournament & Season Administration (Admin)")
add_heading3("Step 1: Season & Rating Class Rule Configuration")
doc.add_paragraph("1. Log into the Admin Dashboard -> Navigate to 'Tournament Management'.\n"
                  "2. Click 'Create Season', enter Season Name, Start Date, and End Date.\n"
                  "3. Under 'Class Rules', define rating brackets (Class 1 to Class 5) by setting min/max ratings and prize ranges.")

add_heading3("Step 2: Scheduling Race Meetings & Races")
doc.add_paragraph("1. Select an active season and click 'Add Race Meeting'. Specify event date, track venue, and total budget.\n"
                  "2. Inside a Race Meeting, click 'Add Race' to configure distance (meters), track type (Turf/Dirt), start time, and registration deadline.")

add_heading3("Step 3: Managing Participant Registrations & Racecards")
doc.add_paragraph("1. Navigate to 'Registration Processing'. Review pending meeting registrations for Owners, Jockeys, and Horses.\n"
                  "2. Click 'Approve' or 'Reject' for each request.\n"
                  "3. Navigate to 'Race Card Approval'. Once declaration closes, approve declared entries. The system automatically assigns starting gate numbers (1 to 14) randomly.")

add_heading3("Step 4: Live Stream Setup & Result Finalization")
doc.add_paragraph("1. Go to 'Live Settings' -> Input YouTube Live URL for active races.\n"
                  "2. After referee submission, navigate to 'Publish Results' to verify finish times and publish official results globally.")

add_heading2("3. Workflow 2: Horse Owner Management")
add_heading3("Step 1: Managing Horse Profiles & Retirement")
doc.add_paragraph("1. Log into Owner Dashboard -> Navigate to 'My Horses'.\n"
                  "2. Click 'Add Horse' to enter horse name, breed, gender, and date of birth.\n"
                  "3. To retire a horse, click 'Request Retirement', enter reason, and submit for Admin approval.")

add_heading3("Step 2: Registering for Meetings & Inviting Jockeys")
doc.add_paragraph("1. Navigate to 'Race Meetings' -> Select upcoming meeting -> Click 'Register Horse'.\n"
                  "2. Once approved by Admin, browse scheduled races and click 'Find Jockey'.\n"
                  "3. Select an available jockey and click 'Send Race Invitation'.")

add_heading2("4. Workflow 3: Jockey Participation & Schedule")
doc.add_paragraph("1. Log into Jockey Dashboard -> Update profile weight and biography.\n"
                  "2. Navigate to 'Meeting Availability' -> Register for upcoming Race Meetings.\n"
                  "3. Go to 'Invitations' -> View pending race invitations from owners -> Click 'Accept' or 'Reject'.\n"
                  "4. Access 'My Race Calendar' to view assigned gate numbers, horse stats, and carried weight calculations.")

add_heading2("5. Workflow 4: Race Officiating (Referee)")
doc.add_paragraph("1. Log into Referee Hub -> Select assigned race.\n"
                  "2. Conduct 'Pre-Race Check': Verify jockey actual weight and horse identity. System calculates carried weight based on handicap rules.\n"
                  "3. Monitor live race progress. If rule infractions occur (e.g. lane cutting), click 'Report Incident' to log violation and penalty.\n"
                  "4. After race completion, click 'Enter Results' -> Enter finish times and positions for all horses -> Submit steward's report.")

add_heading2("6. Workflow 5: Spectator Engagement & AI Assistant")
doc.add_paragraph("1. Visit Landing Page -> Browse public fixtures, racecards, and past results.\n"
                  "2. Click 'Watch Live' to view embedded YouTube Live stream and participate in real-time spectator chat box.\n"
                  "3. Click 'AI Assistant' bubble -> Ask questions about horse stats, jockey win rates, or request AI predictions for upcoming race winners.")

# Save modified document
doc.save(doc_path)
print("Saved Template3_Final Release Document.docx successfully!")
